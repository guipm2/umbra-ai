import io
import logging

import pypdf
import docx
from agno.agent import Agent
from agno.models.openai import OpenAIChat
from agno.knowledge.document import Document
from fastapi import UploadFile

from knowledge_base import get_knowledge_base

logger = logging.getLogger(__name__)

# Allowed file extensions and their magic-byte signatures
ALLOWED_TYPES = {
    "pdf": b"%PDF",
    "docx": b"PK",      # OOXML is a ZIP
    "doc": b"\xd0\xcf",  # OLE2 compound document
    "txt": None,          # No magic bytes for plain text
    "md": None,
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def get_brain_agent():
    kb = get_knowledge_base()
    if not kb:
        return None

    return Agent(
        model=OpenAIChat(id="gpt-4o"),
        knowledge=kb,
        search_knowledge=True,
        description="Você é o Cérebro da empresa. Você tem acesso a todos os documentos internos.",
        instructions=[
            "Sempre pesquise sua base de conhecimento primeiro.",
            "Se a resposta for encontrada nos documentos, cite a fonte.",
            "Se não encontrar, use seu conhecimento geral mas mencione isso.",
        ],
    )


def _validate_file(filename: str, file_bytes: bytes) -> str:
    """Validate file extension and magic bytes. Returns the file type or raises ValueError."""
    if not filename or "." not in filename:
        raise ValueError("Nome de arquivo inválido.")

    file_type = filename.rsplit(".", 1)[-1].lower()

    if file_type not in ALLOWED_TYPES:
        raise ValueError(f"Tipo de arquivo não suportado: .{file_type}")

    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(f"Arquivo excede o limite de {MAX_FILE_SIZE // (1024*1024)}MB.")

    # Validate magic bytes for binary formats
    expected_magic = ALLOWED_TYPES[file_type]
    if expected_magic and not file_bytes[:len(expected_magic)].startswith(expected_magic):
        raise ValueError(f"Conteúdo do arquivo não corresponde à extensão .{file_type}")

    return file_type


async def process_document(file: UploadFile, user_id: str):
    """Extracts text from file and loads it into Knowledge Base."""
    filename = file.filename or "unknown"
    file_bytes = await file.read()

    try:
        file_type = _validate_file(filename, file_bytes)
    except ValueError as e:
        return {"status": "error", "message": str(e)}

    content = ""

    try:
        if file_type == "pdf":
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"

        elif file_type in ("doc", "docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                content += para.text + "\n"

        elif file_type in ("txt", "md"):
            content = file_bytes.decode("utf-8")

        if not content.strip():
            return {"status": "error", "message": "Nenhum texto extraído do documento."}

        kb = get_knowledge_base()
        if kb:
            doc = Document(
                content=content,
                meta_data={
                    "source": filename,
                    "user_id": user_id,
                    "type": file_type,
                },
            )
            kb.vector_db.upsert(content_hash=filename, documents=[doc])
            return {"status": "success", "message": f"Documento processado com sucesso: {filename}"}

        return {"status": "error", "message": "Base de Conhecimento indisponível."}

    except Exception:
        logger.exception("Error processing document: %s", filename)
        return {"status": "error", "message": "Erro ao processar documento."}
